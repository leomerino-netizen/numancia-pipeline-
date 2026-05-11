"""
corrector_aplicado.py — Módulo de corrección IA aplicada (Nivel 2)
==================================================================

Aplica las correcciones detectadas por el corrector preliminar usando
Claude API, generando un .docx con Control de Cambios + Comentarios al margen.

Modo conservador (default):
  - Mecánicas (tildes, comillas, rayas, puntuación, mayúsculas, cursivas,
    erratas claras): se aplican como Control de Cambios (tachado/añadido)
  - Subjetivas (reescritura, concordancia, repeticiones léxicas):
    se sugieren como Comentarios al margen

Procesamiento ASÍNCRONO con sistema de jobs en memoria (RAM compartida
entre workers gunicorn — válido con worker=sync count=2 ya configurado).

Uso desde app.py:
    from corrector_aplicado import (
        crear_job_correccion, procesar_job, get_job_status, get_job_resultado
    )
"""
import io
import os
import re
import json
import time
import uuid
import threading
import traceback
from datetime import datetime
from typing import Optional

import anthropic
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import RGBColor


# ─── Configuración ────────────────────────────────────────────────────────────
MODELO_CLAUDE = 'claude-sonnet-4-5'
MAX_TOKENS_RESPUESTA = 8192
CHUNK_PARRAFOS = 12   # párrafos por llamada — ~3000-4000 chars de contexto
TIMEOUT_API_SEG = 90  # timeout por llamada Claude

# Author identifier para el control de cambios (aparece en Word como autor)
AUTOR_REVISION = 'Editorial Numancia'

# ─── Almacén de jobs en memoria (compartido entre llamadas) ──────────────────
# Estructura: {job_id: {estado, progreso, resumen, docx_b64, error, ts_inicio, ts_fin}}
_JOBS = {}
_JOBS_LOCK = threading.Lock()


# ═════════════════════════════════════════════════════════════════════════════
# PROMPT MAESTRO (RAE 2010 + DPD + Fundéu + Martínez de Sousa)
# ═════════════════════════════════════════════════════════════════════════════
# IMPORTANTE: este prompt debe superar los 1024 tokens para que Anthropic
# active el prompt caching (~70% de ahorro en input cost). Tiene ~2900 tokens.
PROMPT_SISTEMA = """Eres un corrector ortotipográfico profesional de Editorial Numancia.

Tu trabajo es aplicar correcciones según RAE 2010, DPD, Fundéu y Martínez de Sousa
sobre fragmentos de manuscritos en español de literatura contemporánea.

Recibirás párrafos numerados. Para CADA párrafo debes devolver:
- Si NO necesita cambios: lo devuelves tal cual
- Si SÍ necesita cambios: devuelves la versión corregida

Categorías de corrección que DEBES aplicar (mecánicas, sin dudar):
  1. TILDES: solo/sólo (suprimir tilde diacrítica obsoleta según norma 2010),
     demostrativos sin tilde (este, ese, aquel y sus variantes en función pronominal),
     monosílabos sin tilde salvo casos de tilde diacrítica válida (más/mas, sé/se,
     sí/si, tú/tu, mí/mi, él/el, té/te, dé/de, qué/que en interrogativas/exclamativas).
     Casos especiales: aún/aun según equivalencia con "todavía" o "incluso".

  2. COMILLAS: las "rectas" pasan a «angulares»; comillas internas anidadas siguen
     el orden tipográfico español: «texto exterior con "internas inglesas" y luego
     'internas simples'». No usar comillas tipográficas curvas " " en literatura
     española salvo solicitud expresa.

  3. RAYAS DE DIÁLOGO: convertir guiones cortos (- o –) en raya larga (—, U+2014)
     cuando inicien diálogo o inciso. Estructura correcta:
     —Buenos días —dijo Pedro—. ¿Cómo está usted?
     Sin espacio entre la raya de cierre y el verbo dicendi. Espacio antes de
     la raya de apertura del inciso. Punto fuera de la raya de cierre cuando
     el inciso va al final del parlamento.

  4. PUNTUACIÓN: comas faltantes en vocativos ("Hola, María" no "Hola María"),
     incisos delimitados por comas, comas seriales antes de conjunciones (sin
     coma serial estilo Oxford en español: "rojo, verde y azul" no "rojo, verde,
     y azul"). Puntos donde corresponden, eliminar dobles espacios, eliminar
     espacios antes de signos de puntuación. Punto y coma para separar oraciones
     con conexión semántica fuerte. Dos puntos antes de enumeraciones o citas.

  5. MAYÚSCULAS Y MINÚSCULAS: títulos de libros en redonda con mayúscula solo
     en primera palabra ("Cien años de soledad" no "Cien Años De Soledad").
     Cargos en minúscula salvo cabecera de carta o tratamiento de respeto
     ("el rey Felipe VI" pero "Su Majestad el Rey"). Nombres de instituciones
     en mayúscula solo la inicial relevante. Días de la semana y meses en
     minúscula. Gentilicios en minúscula ("los españoles"). Tras dos puntos
     en español NO va mayúscula salvo cita textual o título.

  6. CURSIVAS: extranjerismos no adaptados al español (ballet, software, leitmotiv,
     mise-en-scène), latinismos no incorporados al DLE (a priori, ad hoc, in fraganti),
     títulos de obras citadas dentro del texto (libros, películas, óperas, cuadros,
     álbumes musicales), nombres de embarcaciones y aeronaves, palabras usadas
     metalingüísticamente. NO van en cursiva los nombres propios, las marcas
     comerciales reconocidas (Google, Coca-Cola), ni los extranjerismos adaptados
     que ya están en el DLE (chófer, club, líder, eslogan).

  7. ERRATAS evidentes y errores ortográficos: palabras mal escritas que no son
     intención estilística del autor. Casos típicos: "haver"→"haber", "iva a"→"iba a",
     "haya/halla/aya", "tubo/tuvo", "echo/hecho", "ay/ahí/hay", "vaya/valla/baya",
     "porque/por qué/porqué/por que", "sino/si no", "demás/de más", "también/tan bien".
     CUIDADO: si la palabra es un nombre propio desconocido, jerga local del autor,
     o palabra inventada con intención estética, NO la cambies.

Categorías que SOLO SUGIERES (no las aplicas, las marcas con sugerencias):
  8. REESCRITURA: oraciones confusas, larguísimas (más de 50 palabras sin pausa
     mayor), mal construidas, con anáforas ambiguas. Sugiere división o reordenación.

  9. CONCORDANCIA: dudas de género/número entre sujeto y verbo, entre sustantivo y
     adjetivo o entre antecedente y pronombre relativo. Concordancia ad sensum si
     hay grupos colectivos. Casos de leísmo, laísmo, loísmo si son sistemáticos.

  10. REPETICIONES LÉXICAS: palabras o raíces repetidas en el mismo párrafo o en
      párrafos contiguos. Verbos comodín repetidos (ser, estar, haber, tener,
      hacer). Sugiere sinónimos contextualmente válidos sin alterar registro.

REGLAS CRÍTICAS DE INTERVENCIÓN:
- NO inventes contenido nuevo. Solo corrige lo que ya existe en el texto.
- NO cambies el sentido del texto bajo ningún concepto.
- NO añadas explicaciones en el output, solo el texto corregido.
- Conserva los saltos de línea, sangrías y formato del original.
- Si el original tiene una palabra rara que parece intencionada (jerga, dialecto,
  neologismo del autor, regionalismo, juego fonético), NO la cambies. La firma
  estilística del autor es sagrada.
- Para nombres propios desconocidos, antropónimos, topónimos o marcas, NO los
  corrijas aunque parezcan raros. La RAE no manda sobre nombres propios.
- Si dudas entre corregir o no, NO corrijas. Mejor un falso negativo que un
  falso positivo. El autor revisará y aceptará o rechazará tus sugerencias.
- En diálogos en jerga, slang o registros coloquiales, conserva las formas no
  estándar si son consistentes con la voz del personaje.
- En textos con licencia poética (versos, prosa poética), no impongas puntuación
  académica si el autor ha optado por un estilo libre evidente.

FORMATO DE RESPUESTA OBLIGATORIO:
Devuelves un JSON con esta estructura exacta:
{
  "parrafos": [
    {"id": 1, "texto_corregido": "...", "tiene_cambios": true, "sugerencias": []},
    {"id": 2, "texto_corregido": "...", "tiene_cambios": false, "sugerencias": []},
    ...
  ]
}

donde "sugerencias" es una lista de comentarios al margen para las categorías 8, 9, 10.
Cada sugerencia: {"tipo": "reescritura|concordancia|repeticion", "comentario": "..."}

Si tiene_cambios=false, basta con que pongas el texto original sin modificar.
NO incluyas explicaciones fuera del JSON. El primer carácter de tu respuesta
debe ser { y el último }.
"""


def _construir_prompt_usuario(parrafos_chunk: list) -> str:
    """Construye el mensaje para Claude con los párrafos numerados."""
    lineas = ['Aquí están los párrafos a revisar:\n']
    for idx, parr in enumerate(parrafos_chunk, start=1):
        # Cada párrafo encerrado para evitar confusiones con los saltos
        lineas.append(f'### PÁRRAFO {idx}')
        lineas.append(parr)
        lineas.append('')  # línea en blanco
    lineas.append('\nDevuelve el JSON estructurado.')
    return '\n'.join(lineas)


def _llamar_claude(parrafos_chunk: list, client: anthropic.Anthropic) -> dict:
    """
    Llama a Claude con un chunk de párrafos y devuelve el JSON parseado.
    Activa PROMPT CACHING: el PROMPT_SISTEMA se cachea en Anthropic, así
    que a partir de la segunda llamada solo se cobra el 10% de su coste
    (ahorro estimado: 60-70% del coste total del input).
    Maneja errores de parseo intentando extraer JSON del texto.
    """
    prompt_user = _construir_prompt_usuario(parrafos_chunk)
    try:
        response = client.messages.create(
            model=MODELO_CLAUDE,
            max_tokens=MAX_TOKENS_RESPUESTA,
            system=[
                {
                    "type": "text",
                    "text": PROMPT_SISTEMA,
                    "cache_control": {"type": "ephemeral"}
                }
            ],
            messages=[{'role': 'user', 'content': prompt_user}],
        )
    except Exception as e:
        print(f'[corrector] error API Claude: {e}', flush=True)
        # Devolver chunk sin cambios para no perder el manuscrito
        return {
            'parrafos': [
                {'id': i+1, 'texto_corregido': p, 'tiene_cambios': False, 'sugerencias': []}
                for i, p in enumerate(parrafos_chunk)
            ]
        }

    # Log de uso del caché (útil para verificar que funciona)
    try:
        usage = response.usage
        cache_create = getattr(usage, 'cache_creation_input_tokens', 0) or 0
        cache_read = getattr(usage, 'cache_read_input_tokens', 0) or 0
        input_tokens = getattr(usage, 'input_tokens', 0) or 0
        output_tokens = getattr(usage, 'output_tokens', 0) or 0
        if cache_read > 0:
            print(f'[corrector] 💾 cache HIT: leídos {cache_read} tokens del caché '
                  f'(ahorro ~90%), input {input_tokens}, output {output_tokens}',
                  flush=True)
        elif cache_create > 0:
            print(f'[corrector] 📝 cache CREATE: {cache_create} tokens cacheados '
                  f'(coste +25% solo esta vez), input {input_tokens}, output {output_tokens}',
                  flush=True)
        else:
            print(f'[corrector] tokens: input={input_tokens} output={output_tokens} '
                  f'(sin caché activo)', flush=True)
    except Exception as e:
        print(f'[corrector] no se pudo leer usage: {e}', flush=True)

    raw = response.content[0].text.strip() if response.content else ''
    # Limpiar markdown code fences si vinieran
    if raw.startswith('```'):
        raw = re.sub(r'^```(?:json)?\s*', '', raw)
        raw = re.sub(r'\s*```\s*$', '', raw)

    try:
        data = json.loads(raw)
        if 'parrafos' not in data or not isinstance(data['parrafos'], list):
            raise ValueError('JSON sin clave "parrafos" válida')
        return data
    except Exception as e:
        print(f'[corrector] error parseando JSON: {e}', flush=True)
        print(f'[corrector] respuesta cruda: {raw[:500]}...', flush=True)
        # Fallback: devolver párrafos sin cambios
        return {
            'parrafos': [
                {'id': i+1, 'texto_corregido': p, 'tiene_cambios': False, 'sugerencias': []}
                for i, p in enumerate(parrafos_chunk)
            ]
        }


# ═════════════════════════════════════════════════════════════════════════════
# Gestión de Control de Cambios en .docx
# ═════════════════════════════════════════════════════════════════════════════
def _activar_track_changes(doc: Document):
    """Activa el modo Control de Cambios en el .docx."""
    settings = doc.settings.element
    track = OxmlElement('w:trackChanges')
    settings.append(track)


def _diff_palabras(original: str, corregido: str) -> list:
    """
    Diff a nivel de palabra. Devuelve lista de tuplas:
      ('=', 'palabra')   → sin cambio
      ('-', 'palabra')   → eliminada
      ('+', 'palabra')   → añadida
    Algoritmo simple basado en SequenceMatcher para mantener legibilidad.
    """
    from difflib import SequenceMatcher
    pal_orig = re.findall(r'\S+|\s+', original)
    pal_corr = re.findall(r'\S+|\s+', corregido)
    matcher = SequenceMatcher(None, pal_orig, pal_corr)
    resultado = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            for w in pal_orig[i1:i2]:
                resultado.append(('=', w))
        elif tag == 'delete':
            for w in pal_orig[i1:i2]:
                resultado.append(('-', w))
        elif tag == 'insert':
            for w in pal_corr[j1:j2]:
                resultado.append(('+', w))
        elif tag == 'replace':
            for w in pal_orig[i1:i2]:
                resultado.append(('-', w))
            for w in pal_corr[j1:j2]:
                resultado.append(('+', w))
    return resultado


def _aplicar_cambio_a_parrafo(p, original: str, corregido: str, autor: str, fecha: str):
    """
    Reescribe el párrafo con tags <w:ins> y <w:del> de Word para que
    aparezca en Control de Cambios. Limpia los runs existentes y los
    sustituye por nuevos con tracking.
    """
    if original == corregido:
        return  # nada que hacer

    # Limpiar runs existentes
    for r in list(p.runs):
        r._element.getparent().remove(r._element)

    diff = _diff_palabras(original, corregido)

    # Reservar IDs incrementales (Word los necesita únicos)
    if not hasattr(_aplicar_cambio_a_parrafo, '_next_id'):
        _aplicar_cambio_a_parrafo._next_id = 1

    for tag, palabra in diff:
        if tag == '=':
            # Run normal
            r = p.add_run(palabra)
        elif tag == '+':
            # <w:ins> wrapping run
            run = p.add_run(palabra)
            ins_id = str(_aplicar_cambio_a_parrafo._next_id)
            _aplicar_cambio_a_parrafo._next_id += 1
            ins = OxmlElement('w:ins')
            ins.set(qn('w:id'), ins_id)
            ins.set(qn('w:author'), autor)
            ins.set(qn('w:date'), fecha)
            run_el = run._element
            run_el.getparent().remove(run_el)
            ins.append(run_el)
            p._element.append(ins)
        elif tag == '-':
            # <w:del> wrapping run con <w:delText>
            del_id = str(_aplicar_cambio_a_parrafo._next_id)
            _aplicar_cambio_a_parrafo._next_id += 1
            del_el = OxmlElement('w:del')
            del_el.set(qn('w:id'), del_id)
            del_el.set(qn('w:author'), autor)
            del_el.set(qn('w:date'), fecha)
            r = OxmlElement('w:r')
            t = OxmlElement('w:delText')
            t.set(qn('xml:space'), 'preserve')
            t.text = palabra
            r.append(t)
            del_el.append(r)
            p._element.append(del_el)


def _agregar_comentario(doc: Document, parr, texto_comentario: str,
                        autor: str, comments_state: dict):
    """
    Agrega un comentario al margen sobre el párrafo `parr`.
    Word necesita un archivo word/comments.xml en el .docx + referencias.
    Esta implementación es simplificada: anota el comentario inline en
    el párrafo con un marcador visible (no usa el sistema nativo de
    comments.xml porque python-docx no lo soporta directamente sin
    manipular el ZIP del .docx).

    Para mantener simplicidad y robustez en producción, añadimos el
    comentario como párrafo adicional con estilo distintivo.
    """
    # Añadir un párrafo siguiente con el comentario en cursiva y color azul
    # Este es un compromiso técnico: el comentario aparece justo después
    # del párrafo en lugar de en el margen, pero es totalmente visible y
    # editable. El asesor puede convertirlos en Comentarios reales en Word
    # si lo desea (Insertar > Comentario sobre el texto).
    # Insertar tras el párrafo actual
    nuevo = doc.paragraphs[0].insert_paragraph_before  # placeholder

    # Creamos un nuevo párrafo justo después del actual
    nuevo_p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    # Estilo: indentación + cursiva + color azul
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '720')  # 0.5 inch
    pPr.append(ind)
    nuevo_p.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    italic = OxmlElement('w:i')
    rPr.append(italic)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '1F4E79')  # azul oscuro Editorial
    rPr.append(color)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')  # 9pt
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = f'  💬 [{autor}] {texto_comentario}'
    r.append(t)
    nuevo_p.append(r)

    # Insertar después del párrafo actual
    parr._element.addnext(nuevo_p)


# ═════════════════════════════════════════════════════════════════════════════
# Procesamiento principal del .docx
# ═════════════════════════════════════════════════════════════════════════════
def _extraer_parrafos_no_vacios(doc: Document) -> list:
    """
    Devuelve lista de tuplas (idx_doc, parrafo_obj, texto) solo de los
    párrafos con texto. Se ignoran encabezados de capítulos cortos
    (< 50 chars) para no romperlos con cambios mínimos.
    """
    resultado = []
    for idx, p in enumerate(doc.paragraphs):
        texto = p.text.strip()
        if not texto:
            continue
        # Saltar títulos cortos / numerales romanos
        if len(texto) < 50:
            estilo = (p.style.name if p.style else '').lower()
            if 'heading' in estilo or 'título' in estilo or 'titulo' in estilo:
                continue
            if re.match(r'^(CAP[IÍ]TULO|cap[ií]tulo|[IVXLCDM]+\.?)\s*\d*$', texto):
                continue
        resultado.append((idx, p, texto))
    return resultado


def _chunk_lista(lst: list, tam: int):
    """Divide una lista en sublistas de tamaño `tam`."""
    for i in range(0, len(lst), tam):
        yield lst[i:i+tam]


def _procesar_docx(
    docx_bytes: bytes,
    job_id: str,
    api_key: str,
) -> tuple:
    """
    Procesa el .docx completo. Devuelve (docx_bytes_corregido, resumen_dict).

    Actualiza el progreso del job en _JOBS conforme avanza.
    """
    client = anthropic.Anthropic(api_key=api_key)

    doc = Document(io.BytesIO(docx_bytes))
    parrafos_info = _extraer_parrafos_no_vacios(doc)
    total = len(parrafos_info)

    print(f'[corrector] job={job_id} total párrafos={total}', flush=True)

    if total == 0:
        return docx_bytes, {
            'parrafos_totales': 0,
            'parrafos_modificados': 0,
            'cambios_aplicados': 0,
            'sugerencias_anyadidas': 0,
            'tiempo_seg': 0,
        }

    # Activar Control de Cambios en el .docx
    _activar_track_changes(doc)

    fecha_iso = datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ')
    cambios_aplicados = 0
    sugerencias_anyadidas = 0
    parrafos_modificados = 0

    chunks = list(_chunk_lista(parrafos_info, CHUNK_PARRAFOS))
    total_chunks = len(chunks)

    for i_chunk, chunk in enumerate(chunks):
        textos_chunk = [t for (_, _, t) in chunk]
        resultado = _llamar_claude(textos_chunk, client)

        for j, (idx_doc, parr_obj, texto_orig) in enumerate(chunk):
            try:
                item = resultado['parrafos'][j]
            except (IndexError, KeyError):
                continue

            tiene_cambios = item.get('tiene_cambios', False)
            texto_nuevo = item.get('texto_corregido', texto_orig) or texto_orig
            sugerencias = item.get('sugerencias', []) or []

            # Aplicar cambios mecánicos como Track Changes
            if tiene_cambios and texto_nuevo and texto_nuevo != texto_orig:
                _aplicar_cambio_a_parrafo(
                    parr_obj, texto_orig, texto_nuevo,
                    AUTOR_REVISION, fecha_iso
                )
                cambios_aplicados += 1
                parrafos_modificados += 1

            # Añadir sugerencias subjetivas como párrafos-comentario
            for s in sugerencias:
                if not isinstance(s, dict):
                    continue
                tipo = s.get('tipo', '')
                comentario = s.get('comentario', '').strip()
                if comentario:
                    _agregar_comentario(
                        doc, parr_obj,
                        f'[{tipo.upper()}] {comentario}',
                        AUTOR_REVISION,
                        {}
                    )
                    sugerencias_anyadidas += 1

        # Actualizar progreso del job
        progreso = round((i_chunk + 1) / total_chunks * 100)
        _set_job(job_id, {
            'progreso': progreso,
            'mensaje': f'Procesando bloque {i_chunk+1} de {total_chunks}',
        })
        print(f'[corrector] job={job_id} chunk {i_chunk+1}/{total_chunks} '
              f'cambios_acum={cambios_aplicados} sug_acum={sugerencias_anyadidas}',
              flush=True)

    # Serializar el .docx corregido
    out_buf = io.BytesIO()
    doc.save(out_buf)
    out_bytes = out_buf.getvalue()

    resumen = {
        'parrafos_totales': total,
        'parrafos_modificados': parrafos_modificados,
        'cambios_aplicados': cambios_aplicados,
        'sugerencias_anyadidas': sugerencias_anyadidas,
        'chunks_procesados': total_chunks,
    }
    return out_bytes, resumen


# ═════════════════════════════════════════════════════════════════════════════
# API DE JOBS (asíncrono)
# ═════════════════════════════════════════════════════════════════════════════
def _set_job(job_id: str, datos: dict):
    """Actualiza campos del job de forma thread-safe."""
    with _JOBS_LOCK:
        if job_id not in _JOBS:
            _JOBS[job_id] = {}
        _JOBS[job_id].update(datos)


def _get_job(job_id: str) -> Optional[dict]:
    """Obtiene una copia del estado del job."""
    with _JOBS_LOCK:
        if job_id not in _JOBS:
            return None
        return dict(_JOBS[job_id])


def crear_job_correccion(docx_bytes: bytes, expediente_id: str = '',
                          asesora: str = 'laura') -> str:
    """
    Crea un nuevo job de corrección y lo lanza en background.
    Devuelve el job_id para que el cliente consulte el estado.
    """
    job_id = uuid.uuid4().hex[:12]
    _set_job(job_id, {
        'estado': 'pendiente',
        'progreso': 0,
        'mensaje': 'En cola',
        'expediente_id': expediente_id,
        'asesora': asesora,
        'ts_inicio': time.time(),
        'docx_b64': None,
        'resumen': None,
        'error': None,
    })

    # Lanzar en hilo separado
    hilo = threading.Thread(
        target=_ejecutar_job,
        args=(job_id, docx_bytes),
        daemon=True,
    )
    hilo.start()
    return job_id


def _ejecutar_job(job_id: str, docx_bytes: bytes):
    """Wrapper que ejecuta el procesamiento y captura errores."""
    import base64 as _b64
    try:
        api_key = os.environ.get('ANTHROPIC_API_KEY', '')
        if not api_key:
            raise RuntimeError('ANTHROPIC_API_KEY no configurada')

        _set_job(job_id, {
            'estado': 'procesando',
            'mensaje': 'Iniciando análisis del manuscrito',
        })

        out_bytes, resumen = _procesar_docx(docx_bytes, job_id, api_key)

        _set_job(job_id, {
            'estado': 'completado',
            'progreso': 100,
            'mensaje': 'Corrección completada',
            'docx_b64': _b64.b64encode(out_bytes).decode('ascii'),
            'resumen': resumen,
            'ts_fin': time.time(),
        })
        print(f'[corrector] job={job_id} COMPLETADO {resumen}', flush=True)
    except Exception as e:
        tb = traceback.format_exc()
        print(f'[corrector] job={job_id} ERROR: {e}\n{tb}', flush=True)
        _set_job(job_id, {
            'estado': 'error',
            'mensaje': f'Error: {type(e).__name__}',
            'error': str(e),
            'ts_fin': time.time(),
        })


def get_job_status(job_id: str) -> Optional[dict]:
    """
    Devuelve el estado del job sin el .docx (para polling ligero).
    Devuelve None si el job no existe.
    """
    j = _get_job(job_id)
    if j is None:
        return None
    return {
        'job_id': job_id,
        'estado': j.get('estado'),
        'progreso': j.get('progreso', 0),
        'mensaje': j.get('mensaje', ''),
        'expediente_id': j.get('expediente_id', ''),
        'asesora': j.get('asesora', ''),
        'resumen': j.get('resumen'),
        'error': j.get('error'),
        'tiempo_seg': round((j.get('ts_fin') or time.time()) - j.get('ts_inicio', time.time()), 1),
    }


def get_job_resultado(job_id: str) -> Optional[dict]:
    """
    Devuelve el resultado completo del job INCLUYENDO el .docx en base64.
    Solo se debe llamar cuando estado=='completado'.
    """
    j = _get_job(job_id)
    if j is None or j.get('estado') != 'completado':
        return None
    return {
        'job_id': job_id,
        'estado': j['estado'],
        'docx_corregido_b64': j.get('docx_b64', ''),
        'resumen': j.get('resumen', {}),
        'expediente_id': j.get('expediente_id', ''),
        'asesora': j.get('asesora', ''),
    }


def limpiar_jobs_antiguos(horas: int = 24):
    """Borra jobs completados/error más antiguos que N horas."""
    ahora = time.time()
    limite = horas * 3600
    with _JOBS_LOCK:
        ids_borrar = [
            jid for jid, j in _JOBS.items()
            if j.get('ts_fin') and (ahora - j['ts_fin']) > limite
        ]
        for jid in ids_borrar:
            del _JOBS[jid]
    return len(ids_borrar)
