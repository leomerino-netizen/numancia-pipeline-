"""
docx_parser.py — Parser inteligente de manuscritos Word
Detecta: TOC, título, autor, dedicatoria, epígrafe, capítulos, párrafos, diálogos.
"""
import re
from dataclasses import dataclass, field
from typing import List
from docx import Document
from limpiador_manuscrito import (
    limpiar_bloques, normalizar_comillas, normalizar_dialogos,
    normalizar_espacios, _quitar_invisibles
)

@dataclass
class Bloque:
    tipo: str        # cap_titulo|cap_subtitulo|parrafo|dialogo|separador
    texto: str
    html:  str
    primer_parr: bool = False

@dataclass
class Manuscrito:
    titulo:       str = ''
    autor:        str = ''
    anyo:         str = '2025'
    dedicatoria:  List[str] = field(default_factory=list)
    epigrafe:     List[str] = field(default_factory=list)
    epigrafe_autor: str = ''
    bloques:      List[Bloque] = field(default_factory=list)
    notas_autor:  List[str] = field(default_factory=list)


# ── Regexes ───────────────────────────────────────────────────────────────────
_CAP_RE  = re.compile(
    r'^(CAP[IÍ]TULO\s+\w[\w\s]{0,30}|cap[ií]tulo\s+[\w]+)',
    re.IGNORECASE)
_CAP_DASH= re.compile(
    r'^(CAP[IÍ]TULO\s+[\w]+)\s*[—\-–]\s*(.+)$', re.IGNORECASE)
_ROMANO  = re.compile(r'^[IVXLCDM]{1,6}$')
_DED     = ['para ', 'a ', 'dedicado', 'en memoria', 'a la memoria']

# Palabras estructurales que NO son títulos del libro
_NO_TITULOS = {
    'PRÓLOGO','PROLOGO','PRÓLOGUE','PROLOGUE',
    'INTRODUCCIÓN','INTRODUCCION','INTRODUCTION','INTRO',
    'EPÍLOGO','EPILOGO','EPILOGUE',
    'CAPÍTULO','CAPITULO','CHAPTER','CHAP',
    'PREFACIO','PREFACE','PRÓLOGO DEL AUTOR',
    'AGRADECIMIENTOS','ACKNOWLEDGEMENTS','ACKNOWLEDGMENTS',
    'DEDICATORIA','DEDICATION',
    'PRIMERA PARTE','SEGUNDA PARTE','TERCERA PARTE',
    'PARTE PRIMERA','PARTE SEGUNDA','PARTE TERCERA',
    'PARTE I','PARTE II','PARTE III','PARTE IV',
    'NOTA DEL AUTOR','NOTA DE LA AUTORA','NOTA EDITORIAL',
    'ÍNDICE','INDICE','CONTENTS','SUMARIO',
    'NOVELA','POESÍA','POESIA','ENSAYO','RELATO','CUENTO',
    'EPIGRAFE','EPÍGRAFE','EPIGRAPH',
}


def _runs_html(p):
    plano, html = [], []
    for r in p.runs:
        t = r.text
        if not t: continue
        # Limpieza inline básica de cada run
        t = _quitar_invisibles(t)
        plano.append(t)
        if r.bold and r.italic: html.append(f'<b><i>{t}</i></b>')
        elif r.bold:   html.append(f'<b>{t}</b>')
        elif r.italic: html.append(f'<i>{t}</i>')
        else:          html.append(t)
    plano_t = ''.join(plano).strip()
    html_t  = ''.join(html).strip()
    # Normalización editorial completa
    plano_t = normalizar_comillas(plano_t)
    html_t  = normalizar_comillas(html_t)
    plano_t = normalizar_espacios(plano_t)
    html_t  = normalizar_espacios(html_t)
    return plano_t, html_t


def _estilo(p):
    try: return p.style.name.lower()
    except: return 'normal'


def _tiene_page_break(p):
    """
    Detecta si el párrafo contiene un page break explícito (Ctrl+Enter en Word).
    Busca <w:br w:type="page"/> en el XML del párrafo.
    """
    try:
        xml_str = p._p.xml if hasattr(p, '_p') else ''
        return 'w:type="page"' in xml_str or "w:type='page'" in xml_str
    except Exception:
        return False


def _es_cap(texto, estilo):
    if any(x in estilo for x in ['heading','título','title','chapter']): return True
    if _CAP_RE.match(texto) and len(texto) < 100: return True
    if _ROMANO.match(texto.strip()): return True
    return False


def _es_sep(texto):
    return texto.strip() in ('***','* * *','---','—','§','❧','◆','* *','·')


# ── Parser principal ──────────────────────────────────────────────────────────
def parsear_docx(src) -> Manuscrito:
    import io
    if isinstance(src,(bytes,bytearray)):
        doc = Document(io.BytesIO(src))
    else:
        doc = Document(src)

    ms = Manuscrito()

    # Metadatos del core (solo si tienen sentido — no artículos sueltos)
    cp = doc.core_properties
    _ARTICULOS = {'EL','LA','LOS','LAS','UN','UNA','UNOS','UNAS','DE','DEL','EN'}
    if cp.title and len(cp.title.strip()) >= 3 and cp.title.strip().upper() not in _ARTICULOS:
        ms.titulo = cp.title
    if cp.author and len(cp.author.strip()) >= 2:
        ms.autor  = cp.author

    # Recoger todos los párrafos (incluidos vacíos para detectar páginas en blanco)
    parrs = [(i, p) for i, p in enumerate(doc.paragraphs)]

    # ── 1. Detectar y saltar índice/TOC inicial ───────────────────────────────
    # Criterio: bloque inicial donde TODOS son "Body Text" o tienen
    # el patrón "Capítulo X — título" en la primera sección del doc
    inicio = 0
    # Buscar el primer párrafo que parezca el cuerpo real (Normal y no es lista de caps)
    toc_fin = 0
    for idx, (_, p) in enumerate(parrs):
        est = _estilo(p)
        t   = p.text.strip()
        # Si encontramos el título real (mayúsculas cortas, estilo Normal, no cap list)
        # o una línea que claramente es inicio de texto narrativo, paramos el TOC
        if est == 'normal' and t.isupper() and len(t) < 60 and not _CAP_RE.match(t):
            # Podría ser el título de la obra
            toc_fin = idx
            break
        if est == 'normal' and _es_cap(t, est) and idx > 5:
            toc_fin = idx
            break
        if est not in ('body text', 'normal', 'default paragraph font', ''):
            toc_fin = idx
            break
    else:
        toc_fin = 0

    parrs = parrs[toc_fin:]

    # ── 2. Detectar título y autor ────────────────────────────────────────────
    inicio = 0
    # Palabras que NUNCA son un título completo (artículos, preposiciones)
    _NO_TITULO_SOLO = {'EL','LA','LOS','LAS','UN','UNA','UNOS','UNAS','DE','DEL','EN'}
    idx = 0
    while idx < min(len(parrs), 6):
        p = parrs[idx][1]
        t, h = _runs_html(p)
        est  = _estilo(p)
        t_norm = t.strip().upper()
        es_estructural = t_norm in _NO_TITULOS or any(t_norm.startswith(p) for p in _NO_TITULOS)
        es_articulo_solo = t_norm in _NO_TITULO_SOLO
        es_demasiado_corto = len(t.strip()) < 3

        # CASO ESPECIAL: si vemos un artículo solo ("LA", "EL"…) en mayúsculas
        # y el siguiente párrafo también es texto en mayúsculas corto, los unimos
        # como título compuesto: "LA" + "ATALAYA" → "La Atalaya"
        if es_articulo_solo and t.isupper() and idx + 1 < len(parrs):
            t_next, _ = _runs_html(parrs[idx+1][1])
            t_next_norm = t_next.strip().upper()
            if (t_next.isupper() and len(t_next) < 80 and len(t_next.strip()) >= 3
                and t_next_norm not in _NO_TITULOS
                and not any(t_next_norm.startswith(p) for p in _NO_TITULOS)
                and not _es_cap(t_next, _estilo(parrs[idx+1][1]))):
                # Combinar artículo + sustantivo en formato bonito
                titulo_compuesto = f'{t.strip().capitalize()} {t_next.strip().capitalize()}'
                if not ms.titulo: ms.titulo = titulo_compuesto
                inicio = idx + 2
                # ¿Línea siguiente es el autor?
                if idx+2 < len(parrs):
                    t3, _ = _runs_html(parrs[idx+2][1])
                    est3  = _estilo(parrs[idx+2][1])
                    if t3 and not _es_cap(t3, est3) and len(t3) < 80 and not t3.isupper():
                        if not ms.autor: ms.autor = t3
                        inicio = idx + 3
                break

        if not es_estructural and not es_articulo_solo and not es_demasiado_corto and (
            (t.isupper() and len(t) < 80 and not _es_cap(t, est)) or
            any(x in est for x in ['title','titulo','heading 1'])):
            # Capitalizar bonito en lugar de mantener todo en mayúsculas
            if t.isupper() and not any(x in est for x in ['title','titulo','heading 1']):
                titulo_bonito = t.strip().capitalize() if ' ' not in t.strip() else ' '.join(w.capitalize() for w in t.strip().split())
            else:
                titulo_bonito = t
            if not ms.titulo: ms.titulo = titulo_bonito
            inicio = idx + 1
            # Siguiente podría ser autor o subtítulo
            if idx+1 < len(parrs):
                t2, _ = _runs_html(parrs[idx+1][1])
                est2  = _estilo(parrs[idx+1][1])
                if not _es_cap(t2, est2) and len(t2) < 80 and not t2.isupper():
                    if not ms.autor: ms.autor = t2
                    inicio = idx + 2
            break
        idx += 1

    parrs = parrs[inicio:]

    # ── 3. Detectar zona prelims (dedicatoria / epígrafe) antes del cap 1 ────
    primer_cap_idx = None
    for idx, (_, p) in enumerate(parrs):
        t, _ = _runs_html(p)
        est  = _estilo(p)
        if _es_cap(t, est):
            primer_cap_idx = idx
            break

    if primer_cap_idx and primer_cap_idx > 0:
        # Contar párrafos vacíos consecutivos al final del preámbulo
        # para detectar páginas en blanco intencionales antes del cap 1
        vacios_antes_cap = 0
        for j in range(primer_cap_idx - 1, -1, -1):
            t_j, _ = _runs_html(parrs[j][1])
            if not t_j:
                vacios_antes_cap += 1
            else:
                break
        # 15+ párrafos vacíos seguidos = página en blanco intencional antes del cap 1
        # (ej: portada → blanca → cap 1)
        if vacios_antes_cap >= 15:
            ms.bloques.append(Bloque('pagina_blanca', '', ''))
        # Detectar también page break explícito (Ctrl+Enter) en cualquier párrafo del preámbulo
        else:
            for _, p in parrs[:primer_cap_idx]:
                if _tiene_page_break(p):
                    if not ms.bloques or ms.bloques[-1].tipo != 'pagina_blanca':
                        ms.bloques.append(Bloque('pagina_blanca', '', ''))
                    break

        # Procesar dedicatoria/epígrafe del preámbulo
        for _, p in parrs[:primer_cap_idx]:
            t, h = _runs_html(p)
            tl   = t.lower()
            if any(tl.startswith(s) for s in _DED) and len(t) < 200:
                ms.dedicatoria.append(h or t)
            elif t.startswith(('«','"','"')) or t.startswith('—') and len(t)<120:
                ms.epigrafe.append(h or t)
        parrs = parrs[primer_cap_idx:]

    # ── 4. Parsear cuerpo ─────────────────────────────────────────────────────
    i = 0
    primer_tras_cap = False
    parr_vacios_seguidos = 0   # contador de párrafos vacíos consecutivos
    # Umbral alto: solo secuencias muy largas son verdaderas páginas en blanco
    # intencionales (las separaciones cortas entre capítulos no cuentan).
    UMBRAL_PAG_BLANCA  = 15

    while i < len(parrs):
        _, p = parrs[i]
        t, h = _runs_html(p)
        est  = _estilo(p)

        # ¿Tiene un page break explícito de Word (Ctrl+Enter)?
        # Si el párrafo contiene <w:br w:type="page"/> y está vacío o casi,
        # lo registramos como página en blanco intencional del autor.
        if _tiene_page_break(p) and len(t) < 5:
            # Solo emitimos una pagina_blanca si el bloque previo no es ya una
            if not ms.bloques or ms.bloques[-1].tipo != 'pagina_blanca':
                ms.bloques.append(Bloque('pagina_blanca', '', ''))
            i += 1; continue

        if not t:
            parr_vacios_seguidos += 1
            # Si superamos el umbral, registramos UNA página en blanco
            # (no varias seguidas: solo emitimos al pasar el umbral)
            if parr_vacios_seguidos == UMBRAL_PAG_BLANCA:
                ms.bloques.append(Bloque('pagina_blanca', '', ''))
            i += 1; continue
        else:
            parr_vacios_seguidos = 0

        # Separador
        if _es_sep(t):
            ms.bloques.append(Bloque('separador','❧','❧'))
            i += 1; primer_tras_cap = False; continue

        # Capítulo con guion en misma línea: "CAPÍTULO UNO — El timbre"
        m = _CAP_DASH.match(t)
        if m and len(t) < 100:
            num_cap = m.group(1).strip()
            sub_cap = m.group(2).strip()
            ms.bloques.append(Bloque('cap_titulo', num_cap, num_cap))
            ms.bloques.append(Bloque('cap_subtitulo', sub_cap, sub_cap))
            primer_tras_cap = True; i += 1; continue

        # Capítulo solo
        if _es_cap(t, est):
            num_cap = t.strip()
            sub_cap = ''
            # Mirar si la siguiente línea es subtítulo
            if i+1 < len(parrs):
                t2, h2 = _runs_html(parrs[i+1][1])
                est2   = _estilo(parrs[i+1][1])
                if t2 and not _es_cap(t2, est2) and len(t2) < 100 \
                   and not t2.startswith('—') and len(t2.split()) < 14:
                    sub_cap = t2; i += 1
            ms.bloques.append(Bloque('cap_titulo', num_cap, num_cap))
            if sub_cap:
                ms.bloques.append(Bloque('cap_subtitulo', sub_cap, sub_cap))
            primer_tras_cap = True; i += 1; continue

        # Nota del autor
        if (t.startswith('[') and t.endswith(']')) or 'note' in est:
            ms.notas_autor.append(t.strip('[]'))
            i += 1; continue

        # Diálogo
        if t.startswith(('—','\u2014')):
            ms.bloques.append(Bloque('dialogo', t, h or t,
                                      primer_parr=primer_tras_cap))
            primer_tras_cap = False; i += 1; continue

        # Párrafo normal
        ms.bloques.append(Bloque('parrafo', t, h or t,
                                  primer_parr=primer_tras_cap))
        primer_tras_cap = False; i += 1

    # Limpieza editorial final de todos los bloques
    limpiar_bloques(ms.bloques)

    # Para cada bloque tipo 'parrafo' que empiece con guión variado, marcarlo dialogo
    for b in ms.bloques:
        if b.tipo == 'parrafo' and b.texto.lstrip().startswith('—'):
            b.tipo = 'dialogo'

    return ms


if __name__ == '__main__':
    import sys
    ms = parsear_docx(sys.argv[1] if len(sys.argv)>1 else '/mnt/user-data/uploads/Sara.docx')
    print(f"Título:    '{ms.titulo}'")
    print(f"Autor:     '{ms.autor}'")
    print(f"Dedicatoria: {ms.dedicatoria}")
    print(f"Epígrafe:    {ms.epigrafe}")
    print(f"Bloques: {len(ms.bloques)}")
    print("\nPrimeros 15 bloques:")
    for b in ms.bloques[:15]:
        print(f"  [{b.tipo:15}] p={b.primer_parr} | {b.texto[:70]}")
