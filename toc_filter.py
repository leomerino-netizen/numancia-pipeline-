"""
toc_filter.py v2 — Filtro robusto de Tabla de Contenidos para Editorial Numancia.

Detecta y elimina bloques TOC del manuscrito antes del procesamiento downstream
(preview, maqueta, informe, corrector). Diseñado para manuscritos en cualquier
formato: TOC manual, TOC nativo de Word, TOC por estilos, sin TOC.

Cobertura:
  1. TOC manual con marcador explícito (ÍNDICE, CONTENIDO, SUMARIO, TOC...).
  2. TOC manual sin marcador, detectado por leader dots / tabs / underscores.
  3. TOC nativo de Word (campo <w:fldChar> con instrucción TOC).
  4. TOC por estilos nativos de Word (TOC1, TOC2, TOC3...).

Defensas:
  - Safety cap: ningún rango puede superar SAFETY_CAP_RATIO del documento.
  - Whitelist por posición: marcadores en el último POS_WHITELIST_RATIO se
    ignoran (probablemente índice onomástico/temático/glosario).
  - Whitelist por nombre: "Índice onomástico", "Glosario", "Bibliografía" 
    nunca se tratan como TOC.
  - Logging auditable: emite por logger "editorial.toc_filter".
  - Multi-bloque: maneja varios bloques TOC en el mismo documento.

Uso en docx_parser.py:

    from toc_filter import filtrar_indice_manual

    doc = Document(src)
    parrafos = filtrar_indice_manual(list(doc.paragraphs), docx_path=src)
    # resto del pipeline con `parrafos`

Compatible con: python-docx Paragraph, str, dict {'text': '...'}.
"""

from __future__ import annotations

import logging
import re
import unicodedata
from dataclasses import dataclass, field
from typing import Any, List, Optional, Tuple
from zipfile import ZipFile, BadZipFile

logger = logging.getLogger("editorial.toc_filter")


# ============================================================================
# Configuración por defecto
# ============================================================================

SAFETY_CAP_RATIO = 0.35         # Rango máximo permitido (35% del documento)
POS_WHITELIST_RATIO = 0.85      # Tras este punto, marcadores → índice temático
DEFAULT_VENTANA_CIERRE = 5      # Párrafos consecutivos de prosa para cerrar
DEFAULT_UMBRAL_PROSA = 25       # Palabras mínimas para considerar prosa


# ============================================================================
# Patrones
# ============================================================================

_MARCADORES_TOC = {
    "indice",
    "indice general",
    "tabla de contenidos",
    "tabla de contenido",
    "contenido",
    "contenidos",
    "sumario",
    "table of contents",
    "toc",
}

_MARCADORES_NO_TOC = {
    "indice onomastico",
    "indice tematico",
    "indice analitico",
    "indice de nombres",
    "indice de materias",
    "indice de autores",
    "indice de ilustraciones",
    "indice de figuras",
    "indice de tablas",
    "glosario",
    "bibliografia",
    "referencias",
    "referencias bibliograficas",
    "fuentes",
    "notas",
    "agradecimientos",
}

_RE_LEADER = re.compile(
    r"(?:"
    r"\.{3,}"
    r"|\u2026{1,}"
    r"|_{3,}"
    r"|(?:\.\s){2,}\."
    r"|\t{2,}"
    r")"
    r"[\s.\u2026_\t]*"
    r"\s*\d{1,4}(?:[\s,\-–—]+\d{1,4})*\s*$"  # un nº o lista: "12", "12, 45", "12-45, 67"
)

_RE_NUMERACION_TOC = re.compile(r"^\s*\d+(?:\.\d+){1,3}\s+\S")

_RE_CAPITULO_NUM = re.compile(
    r"^\s*cap[ií]tulo\s+\d+\s*[:.\-–—]?\s*", re.IGNORECASE
)

_RE_STYLE_TOC = re.compile(r"^TOC\d*$|^toc\s*\d*$|^Contents\d*$", re.IGNORECASE)


# ============================================================================
# Estructuras de resultado
# ============================================================================

@dataclass
class RangoTOC:
    inicio: int
    fin: int
    motivo: str
    descartado: bool = False
    razon_descarte: str = ""

    @property
    def tamano(self) -> int:
        return self.fin - self.inicio + 1


@dataclass
class ResultadoAnalisis:
    total_parrafos: int
    rangos: List[RangoTOC] = field(default_factory=list)

    @property
    def rangos_aplicados(self) -> List[RangoTOC]:
        return [r for r in self.rangos if not r.descartado]

    @property
    def rangos_descartados(self) -> List[RangoTOC]:
        return [r for r in self.rangos if r.descartado]

    @property
    def parrafos_filtrados(self) -> int:
        return sum(r.tamano for r in self.rangos_aplicados)


# ============================================================================
# Utilidades
# ============================================================================

def _normalizar(s: str) -> str:
    if not s:
        return ""
    s = unicodedata.normalize("NFD", s)
    s = "".join(c for c in s if unicodedata.category(c) != "Mn")
    return re.sub(r"\s+", " ", s.strip().lower())


def _texto(obj: Any) -> str:
    if isinstance(obj, str):
        return obj
    if hasattr(obj, "text"):
        return obj.text or ""
    if isinstance(obj, dict):
        return obj.get("text", "") or ""
    return str(obj)


def _estilo(obj: Any) -> str:
    style = getattr(obj, "style", None)
    if style is None:
        return ""
    name = getattr(style, "name", None) or getattr(style, "style_id", None) or ""
    return str(name)


def _es_marcador_toc(texto: str) -> bool:
    norm = _normalizar(texto)
    if not norm or len(norm) > 35:
        return False
    return norm in _MARCADORES_TOC


def _es_marcador_no_toc(texto: str) -> bool:
    norm = _normalizar(texto)
    if not norm or len(norm) > 60:
        return False
    return norm in _MARCADORES_NO_TOC


def _es_item_toc(texto: str) -> bool:
    if not texto:
        return False
    return bool(_RE_LEADER.search(texto))


def _parece_prosa(texto: str, umbral_palabras: int) -> bool:
    t = (texto or "").strip()
    if not t:
        return False
    if _es_item_toc(t):
        return False
    if _RE_NUMERACION_TOC.match(t):
        return False
    if _RE_CAPITULO_NUM.match(t):
        return False
    letras = [c for c in t if c.isalpha()]
    if letras and sum(1 for c in letras if c.isupper()) / len(letras) > 0.75:
        return False
    return len(t.split()) >= umbral_palabras


def _fusionar_rangos(rangos: List[RangoTOC]) -> List[RangoTOC]:
    if not rangos:
        return []
    ordenados = sorted(rangos, key=lambda r: (r.inicio, r.fin))
    fusionados = [ordenados[0]]
    for r in ordenados[1:]:
        ult = fusionados[-1]
        if r.inicio <= ult.fin + 2:
            motivo = ult.motivo if ult.motivo == r.motivo else f"{ult.motivo}+{r.motivo}"
            fusionados[-1] = RangoTOC(
                inicio=ult.inicio,
                fin=max(ult.fin, r.fin),
                motivo=motivo,
            )
        else:
            fusionados.append(r)
    return fusionados


# ============================================================================
# Detectores
# ============================================================================

def _expandir_rango(
    textos: List[str],
    desde: int,
    ventana_cierre: int,
    umbral_prosa: int,
) -> int:
    n = len(textos)
    fin = desde
    j = desde + 1
    consecutivos_prosa = 0
    while j < n:
        if _parece_prosa(textos[j], umbral_prosa):
            consecutivos_prosa += 1
            if consecutivos_prosa >= ventana_cierre:
                return j - ventana_cierre
        else:
            consecutivos_prosa = 0
            fin = j
        j += 1
    return fin


def _detectar_por_texto(
    parrafos: List[Any],
    ventana_cierre: int,
    umbral_prosa: int,
) -> List[RangoTOC]:
    textos = [_texto(p) for p in parrafos]
    n = len(textos)
    rangos: List[RangoTOC] = []
    i = 0

    while i < n:
        if _es_marcador_toc(textos[i]):
            inicio = i
            fin = _expandir_rango(textos, i, ventana_cierre, umbral_prosa)
            rangos.append(RangoTOC(inicio=inicio, fin=fin, motivo="marcador_texto"))
            i = fin + 1
            continue

        if _es_item_toc(textos[i]):
            ventana_inicio = max(0, i - 5)
            ventana_fin = min(n, i + 20)
            cuenta = sum(
                1 for k in range(ventana_inicio, ventana_fin) if _es_item_toc(textos[k])
            )
            if cuenta >= 3:
                inicio = i
                while inicio > 0 and (
                    _es_item_toc(textos[inicio - 1])
                    or _RE_CAPITULO_NUM.match(textos[inicio - 1] or "")
                    or _RE_NUMERACION_TOC.match(textos[inicio - 1] or "")
                ):
                    inicio -= 1
                fin = _expandir_rango(textos, i, ventana_cierre, umbral_prosa)
                rangos.append(RangoTOC(inicio=inicio, fin=fin, motivo="leader_dots"))
                i = fin + 1
                continue

        i += 1

    return rangos


def _detectar_por_estilos(parrafos: List[Any]) -> List[RangoTOC]:
    rangos: List[RangoTOC] = []
    en_bloque = False
    inicio = 0

    for i, p in enumerate(parrafos):
        estilo = _estilo(p)
        es_toc = bool(_RE_STYLE_TOC.match(estilo)) if estilo else False

        if es_toc and not en_bloque:
            en_bloque = True
            inicio = i
        elif not es_toc and en_bloque:
            rangos.append(RangoTOC(inicio=inicio, fin=i - 1, motivo="estilo_toc"))
            en_bloque = False

    if en_bloque:
        rangos.append(RangoTOC(inicio=inicio, fin=len(parrafos) - 1, motivo="estilo_toc"))

    return rangos


def _detectar_toc_word_nativo(docx_path: str) -> List[Tuple[int, int]]:
    """
    Detecta campos TOC nativos de Word (<w:fldChar> con instrucción TOC).
    Devuelve rangos de índices de párrafos (0-based) ocupados por el TOC.
    """
    try:
        with ZipFile(docx_path) as z:
            xml = z.read("word/document.xml").decode("utf-8", errors="replace")
    except (BadZipFile, KeyError, OSError) as e:
        logger.debug("No se pudo leer XML para detección de TOC nativo: %s", e)
        return []

    body_match = re.search(r"<w:body\b[^>]*>(.*)</w:body>", xml, re.DOTALL)
    if not body_match:
        return []
    body = body_match.group(1)

    rangos: List[Tuple[int, int]] = []
    parrafo_idx = -1
    en_toc = False
    toc_inicio: Optional[int] = None
    esperando_instr = False
    instr_buffer = ""

    token_re = re.compile(
        r"<w:p\b[^>]*?/>"
        r"|<w:p\b[^>]*?>"
        r"|</w:p>"
        r"|<w:fldChar\b[^>]*?/>"
        r"|<w:fldChar\b[^>]*?>.*?</w:fldChar>"
        r"|<w:instrText\b[^>]*?>(.*?)</w:instrText>",
        re.DOTALL,
    )

    for m in token_re.finditer(body):
        tok = m.group(0)

        # Apertura de párrafo (con o sin auto-cierre)
        if tok.startswith("<w:p") and not tok.startswith("<w:pPr") and not tok.startswith("<w:pStyle"):
            if tok.endswith("/>"):
                parrafo_idx += 1
                continue
            if tok.startswith("<w:p>") or tok.startswith("<w:p "):
                parrafo_idx += 1
                continue

        if tok == "</w:p>":
            continue

        if tok.startswith("<w:fldChar"):
            if 'fldCharType="begin"' in tok:
                esperando_instr = True
                instr_buffer = ""
            elif 'fldCharType="end"' in tok and en_toc:
                if toc_inicio is not None:
                    rangos.append((toc_inicio, parrafo_idx))
                en_toc = False
                toc_inicio = None
                esperando_instr = False
            continue

        if tok.startswith("<w:instrText"):
            if esperando_instr:
                instr_inner = m.group(1) or ""
                instr_buffer += instr_inner
                if re.search(r"\bTOC\b", instr_buffer):
                    en_toc = True
                    toc_inicio = parrafo_idx
                    esperando_instr = False
            continue

    return rangos


# ============================================================================
# API pública
# ============================================================================

def analizar(
    parrafos: List[Any],
    docx_path: Optional[str] = None,
    safety_cap: float = SAFETY_CAP_RATIO,
    posicion_whitelist: float = POS_WHITELIST_RATIO,
    ventana_cierre: int = DEFAULT_VENTANA_CIERRE,
    umbral_prosa: int = DEFAULT_UMBRAL_PROSA,
) -> ResultadoAnalisis:
    """
    Analiza el manuscrito y devuelve todos los rangos TOC con su veredicto.
    Útil para auditoría sin filtrar nada.
    """
    n = len(parrafos)
    resultado = ResultadoAnalisis(total_parrafos=n)
    if n == 0:
        return resultado

    candidatos: List[RangoTOC] = []
    candidatos.extend(_detectar_por_texto(parrafos, ventana_cierre, umbral_prosa))
    candidatos.extend(_detectar_por_estilos(parrafos))

    if docx_path:
        for ini, fin in _detectar_toc_word_nativo(docx_path):
            candidatos.append(RangoTOC(inicio=ini, fin=fin, motivo="fldchar_word"))

    candidatos = _fusionar_rangos(candidatos)

    textos = [_texto(p) for p in parrafos]
    for r in candidatos:
        # Defensa 1: safety cap
        if r.tamano / n > safety_cap:
            r.descartado = True
            r.razon_descarte = (
                f"supera safety_cap ({r.tamano}/{n} = {r.tamano/n:.0%} > {safety_cap:.0%})"
            )
            resultado.rangos.append(r)
            continue

        # Defensa 2: whitelist por posición
        if r.inicio / n > posicion_whitelist:
            r.descartado = True
            r.razon_descarte = (
                f"posición tardía ({r.inicio}/{n} = {r.inicio/n:.0%} > "
                f"{posicion_whitelist:.0%}), probable índice temático/glosario"
            )
            resultado.rangos.append(r)
            continue

        # Defensa 3: whitelist por nombre
        # Mirar los primeros 3 párrafos del rango Y los 3 anteriores (el marcador
        # puede haber quedado fuera del rango si se detectó por leader_dots).
        ventana_chequeo = list(range(max(0, r.inicio - 3), r.inicio)) + \
                          list(range(r.inicio, min(r.inicio + 3, n)))
        for k in ventana_chequeo:
            if _es_marcador_no_toc(textos[k]):
                r.descartado = True
                r.razon_descarte = f"marcador NO-TOC en párrafo {k}: «{textos[k][:50]}»"
                break

        resultado.rangos.append(r)

    if logger.isEnabledFor(logging.INFO):
        for r in resultado.rangos:
            if r.descartado:
                logger.info(
                    "TOC descartado [%d–%d] motivo=%s: %s",
                    r.inicio, r.fin, r.motivo, r.razon_descarte,
                )
            else:
                logger.info(
                    "TOC aplicado [%d–%d] (%d párrafos) motivo=%s",
                    r.inicio, r.fin, r.tamano, r.motivo,
                )

    return resultado


def filtrar_indice_manual(
    parrafos: List[Any],
    docx_path: Optional[str] = None,
    safety_cap: float = SAFETY_CAP_RATIO,
    posicion_whitelist: float = POS_WHITELIST_RATIO,
    ventana_cierre: int = DEFAULT_VENTANA_CIERRE,
    umbral_prosa: int = DEFAULT_UMBRAL_PROSA,
) -> List[Any]:
    """
    Devuelve la lista de párrafos sin el bloque TOC.

    Args:
        parrafos: lista de párrafos (python-docx, str o dict).
        docx_path: ruta al .docx original (opcional, habilita detección XML).
        safety_cap: fracción máxima del documento que puede ser TOC.
        posicion_whitelist: tras esta fracción, marcadores se ignoran.
        ventana_cierre: párrafos de prosa consecutivos para cerrar el bloque.
        umbral_prosa: palabras mínimas para considerar prosa.
    """
    resultado = analizar(
        parrafos,
        docx_path=docx_path,
        safety_cap=safety_cap,
        posicion_whitelist=posicion_whitelist,
        ventana_cierre=ventana_cierre,
        umbral_prosa=umbral_prosa,
    )
    if not resultado.rangos_aplicados:
        return list(parrafos)

    a_excluir = set()
    for r in resultado.rangos_aplicados:
        for k in range(r.inicio, r.fin + 1):
            a_excluir.add(k)

    return [p for k, p in enumerate(parrafos) if k not in a_excluir]


# ============================================================================
# CLI para debug y auditoría
# ============================================================================

if __name__ == "__main__":
    import sys

    logging.basicConfig(
        level=logging.INFO,
        format="%(levelname)s [%(name)s] %(message)s",
    )

    if len(sys.argv) < 2:
        print("Uso: python toc_filter.py <archivo.docx> [<archivo2.docx> ...]")
        sys.exit(1)

    try:
        from docx import Document
    except ImportError:
        print("ERROR: instala python-docx → pip install python-docx")
        sys.exit(1)

    for path in sys.argv[1:]:
        try:
            doc = Document(path)
        except Exception as e:
            print(f"\n[ERROR] No se pudo abrir {path}: {e}")
            continue

        parrafos = list(doc.paragraphs)
        resultado = analizar(parrafos, docx_path=path)

        print(f"\n{'='*72}")
        print(f"ARCHIVO: {path}")
        print(f"Total de párrafos: {resultado.total_parrafos}")
        print(f"Rangos TOC detectados: {len(resultado.rangos)}")
        print(f"  Aplicados:    {len(resultado.rangos_aplicados)}")
        print(f"  Descartados:  {len(resultado.rangos_descartados)}")
        print(f"Párrafos a filtrar: {resultado.parrafos_filtrados}")
        print("="*72)

        for r in resultado.rangos:
            estado = "DESCARTADO" if r.descartado else "APLICADO  "
            print(f"\n[{estado}] [{r.inicio}–{r.fin}] · {r.tamano} párrafos · {r.motivo}")
            if r.descartado:
                print(f"  Razón: {r.razon_descarte}")
            for k in range(r.inicio, min(r.fin + 1, r.inicio + 4)):
                print(f"  [{k}] {parrafos[k].text[:88]}")
            if r.tamano > 4:
                print(f"  ... ({r.tamano - 4} más)")
                print(f"  [{r.fin}] {parrafos[r.fin].text[:88]}")

        limpios = filtrar_indice_manual(parrafos, docx_path=path)
        print(f"\nResultado: {len(parrafos)} → {len(limpios)} párrafos "
              f"(-{len(parrafos) - len(limpios)})")
